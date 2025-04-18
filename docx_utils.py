# docx_utils.py
from math import ceil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import (
    WD_COLOR_INDEX,
)  # pour la surbrillance de texte (limité aux couleurs prédéfinies)


def trim_value(value: str) -> str:
    """
    Supprime les espaces en début/fin et les guillemets éventuels d'une valeur lue du CSV.

    :param value: La valeur à nettoyer
    :return: La valeur nettoyée
    """
    if isinstance(value, str):
        return value.strip().strip('"')
    elif isinstance(value, list):
        # Appliquer le trim à chacun des éléments et les joindre par une virgule
        return ", ".join(
            [
                v.strip().strip('"') if isinstance(v, str) else str(v)
                for v in value
            ]
        )
    else:
        return str(value)


def compute_duration(start: str, end: str) -> str:
    """
    Calcule la durée (en années) entre une date de début et de fin si possible.

    :param start: Date de début (ex: "2019")
    :param end: Date de fin (ex: "2021")
    :return: Chaîne de caractères du type ", 2 ans" ou une chaîne vide en cas d'impossibilité
    """
    try:
        start_year = int(start)
        end_year = int(end)
        duration = end_year - start_year
        return f", {duration} ans" if duration > 0 else ""
    except ValueError:
        return ""


def add_skills_table(
    document: Document, skills: list, num_cols: int = 2, font_size: Pt = Pt(12)
) -> None:
    """
    Ajoute un tableau des compétences au document.

    :param document: Instance du document (Document)
    :param skills: Liste des dictionnaires de compétences
    :param num_cols: Nombre de colonnes dans le tableau (par défaut 2)
    :param font_size: Taille de police pour chaque cellule
    """
    num_skills = len(skills)
    if num_skills == 0:
        return
    num_rows = ceil(num_skills / num_cols)
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"
    for i, skill in enumerate(skills):
        row_index = i // num_cols
        col_index = i % num_cols
        cell = table.cell(row_index, col_index)
        # Format : "Titre: Description"
        text = f"{trim_value(skill.get('title', ''))}: {trim_value(skill.get('description', ''))}"
        para = cell.paragraphs[0]
        para.text = text
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.size = font_size


def add_custom_heading(
    document: Document,
    text: str,
    level: int = 1,
    font_size: Pt = Pt(24),
    color: RGBColor = RGBColor(0x00, 0x00, 0x00),
    alignment: int = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    """
    Ajoute un titre personnalisé au document avec des styles définis.

    :param document: Instance du document (Document)
    :param text: Texte du titre
    :param level: Niveau du titre (par défaut 1)
    :param font_size: Taille de la police (par défaut 24 pts)
    :param color: Couleur de la police (par défaut rouge)
    :param alignment: Alignement du titre (par défaut centré)
    """
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = font_size
        run.font.color.rgb = color
    heading.alignment = alignment


def add_custom_paragraph(
    document: Document,
    text: str,
    alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
    space_after: Pt = Pt(12),
    font_size: Pt = None,
    font_color: RGBColor = None,
    bold: bool = False,
) -> None:
    """
    Ajoute un paragraphe personnalisé avec une mise en forme simple.

    :param document: Instance du document (Document)
    :param text: Texte du paragraphe
    :param alignment: Alignement du paragraphe (par défaut à gauche)
    :param space_after: Espacement après le paragraphe (par défaut 12 pts)
    :param font_size: Taille de la police (optionnel)
    :param font_color: Couleur de la police (optionnel)
    :param bold: Indique si le texte doit être en gras (par défaut False)
    :return: Le paragraphe créé
    """
    p = document.add_paragraph(text)
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    for run in p.runs:
        if font_size is not None:
            run.font.size = font_size
        if font_color is not None:
            run.font.color.rgb = font_color
        if bold:
            run.bold = True
    return p


def add_styled_run(
    paragraph,
    text: str,
    font_size: Pt = None,
    font_color: RGBColor = None,
    bold: bool = False,
    highlight: int = None,
) -> None:
    """
    Ajoute un run stylisé à un paragraphe.

    :param paragraph: Le paragraphe auquel ajouter le run
    :param text: Texte du run
    :param font_size: Taille de la police (optionnel)
    :param font_color: Couleur de la police (optionnel)
    :param bold: Indique si le texte doit être en gras (par défaut False)
    :param highlight: Couleur de surbrillance (utiliser une valeur de WD_COLOR_INDEX, par exemple WD_COLOR_INDEX.YELLOW)
    :return: Le run ajouté
    """
    run = paragraph.add_run(text)
    if font_size is not None:
        run.font.size = font_size
    if font_color is not None:
        run.font.color.rgb = font_color
    if bold:
        run.bold = True
    if highlight is not None:
        run.font.highlight_color = highlight
    return run


def set_paragraph_format(
    paragraph, alignment: int = None, space_after: Pt = None
) -> None:
    """
    Configure l'alignement et l'espacement après d'un paragraphe.

    :param paragraph: Le paragraphe à formater
    :param alignment: Valeur de l'alignement (WD_ALIGN_PARAGRAPH.LEFT, RIGHT, CENTER, JUSTIFY), optionnel
    :param space_after: Espacement après le paragraphe (en Pt), optionnel
    """
    if alignment is not None:
        paragraph.alignment = alignment
    if space_after is not None:
        paragraph.paragraph_format.space_after = space_after
