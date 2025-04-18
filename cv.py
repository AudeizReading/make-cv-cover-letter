import csv
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_utils import (
    add_custom_heading,
    add_custom_paragraph,
    add_styled_run,
    set_paragraph_format,
    trim_value,
    compute_duration,
)


def generate_cv(csv_file, output_file, cv_type: str = "debutant"):
    """
    Lit le CSV de données de CV et génère un document RTF.
    """
    # Dictionnaires pour chaque section
    personal = []
    objectives = []
    experiences = []
    education = []
    skills = []
    title = ""

    # Lecture du fichier CSV du CV
    with open(csv_file, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row = {key: trim_value(value) for key, value in row.items()}
            section = row.get("section", "").lower()
            if section == "personal":
                personal.append(row)
            elif section == "experience":
                experiences.append(row)
            elif section == "education":
                education.append(row)
            elif section == "skills":
                skills.append(row)
            elif section == "title":
                title = row.get("description", "").strip()
            elif section == "objectives":
                objectives.append(row)

    # Créer un document DOCX
    doc = Document()

    # Informations personnelles
    if personal:
        for item in personal:
            # On tente d'afficher une information avec priorité description,
            # content, puis subtitle
            info = (
                item.get("description")
                or item.get("content")
                or item.get("subtitle")
            )
            para_text = (
                f"{item['title']}: {info}"
                if item["title"] == "Tel"
                else f"{info}"
            )
            add_custom_paragraph(
                doc,
                para_text,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(2),
                font_size=Pt(12),
                font_color=RGBColor(0x00, 0x00, 0x00),
            )

    # Titre du CV
    add_custom_heading(
        doc,
        title,
        level=1,
        font_size=Pt(24),
        color=RGBColor(0x00, 0x00, 0x00),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    if cv_type.lower() == "debutant" and objectives:
        objective_text = objectives[0].get("description", "").strip()
        add_custom_paragraph(
            doc,
            objective_text,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(12),
            font_size=Pt(12),
            bold=True,
        )

    # Fonctions internes pour ajouter les sections
    def add_experiences_section(doc, experiences):
        add_custom_heading(
            doc,
            "Expériences professionnelles",
            level=2,
            font_size=Pt(18),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        for exp in experiences:
            start_date = exp.get("start_date", "")
            end_date = exp.get("end_date", "")
            if start_date and end_date:
                duration_str = compute_duration(start_date, end_date)
                date_str = f" ({start_date} - {end_date}{duration_str})"
            elif start_date:
                date_str = f" ({start_date})"
            else:
                date_str = ""
            p = doc.add_paragraph()
            add_styled_run(
                p,
                f"{exp.get('title')} - {exp.get('subtitle')}{date_str}",
                font_size=Pt(14),
                bold=True,
            )
            set_paragraph_format(p, space_after=Pt(4))
            add_custom_paragraph(
                doc,
                exp.get("description"),
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_after=Pt(6),
            )

    def add_education_section(doc, education):
        add_custom_heading(
            doc,
            "Formation",
            level=2,
            font_size=Pt(18),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        for edu in education:
            start_date = edu.get("start_date", "")
            end_date = edu.get("end_date", "")
            if start_date and end_date:
                date_str = f" ({start_date} - {end_date})"
            elif start_date:
                date_str = f" ({start_date})"
            else:
                date_str = ""
            p = doc.add_paragraph()
            add_styled_run(
                p,
                f"{edu.get('title')} - {edu.get('subtitle')}{date_str}",
                font_size=Pt(14),
                bold=True,
            )
            set_paragraph_format(p, space_after=Pt(4))
            add_custom_paragraph(
                doc,
                edu.get("description"),
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_after=Pt(6),
            )

    # Ordre des sections selon le type de CV
    if cv_type.lower() == "accompli" and experiences and education:
        add_experiences_section(doc, experiences)
        add_education_section(doc, education)
    elif experiences and education:
        add_education_section(doc, education)
        add_experiences_section(doc, experiences)

    # Compétences
    if skills:
        add_custom_heading(
            doc,
            "Compétences",
            level=2,
            font_size=Pt(18),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        for skill in skills:
            competence = f"- {skill['title']}: {skill['description']}"
            add_custom_paragraph(
                doc,
                competence,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(4),
                font_size=Pt(12),
            )

    # Sauvegarde du document DOCX
    doc.save(output_file)
    print(f"CV généré et sauvegardé dans {output_file}")


if __name__ == "__main__":
    # Fichiers CSV d'entrée
    import sys

    if len(sys.argv) < 2:
        print("Usage: python cv.py <cv_csv_file> [cv_type]")
        sys.exit(1)
    cv_csv = sys.argv[1]
    cv_type = sys.argv[2] if len(sys.argv) > 2 else "debutant"
    cv_output = "cv-test.docx"
    if os.path.exists(cv_csv):
        generate_cv(cv_csv, cv_output, cv_type=cv_type)
    else:
        print(f"Le fichier {cv_csv} n'existe pas.")
