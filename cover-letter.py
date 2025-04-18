import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_utils import (
    trim_value,
)


# TODO: Il reste pas mal de mise en forme à faire
def generate_docx_cover_letter(csv_file, output_file):
    """
    Lit le fichier CSV de la lettre de motivation et génère un document DOCX.
    """
    # Lecture du CSV
    data = {}
    with open(csv_file, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row = {key: trim_value(value) for key, value in row.items()}
            data[row.get("title", "")] = row.get("content", "")

    # Création du document
    doc = Document()

    # Coordonnées de l'expéditeur (en haut à gauche)
    if "Expéditeur" in data:
        p_expediteur = doc.add_paragraph(data["Expéditeur"])
        p_expediteur.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_expediteur.paragraph_format.space_after = Pt(12)

    # Laisser quelques lignes d'espace
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Coordonnées du destinataire (affichées ensuite, tjs alignées à gauche)
    if "Destinataire" in data:
        p_destinataire = doc.add_paragraph(data["Destinataire"])
        p_destinataire.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_destinataire.paragraph_format.space_after = Pt(12)

    # Laisser un espace avant la date
    doc.add_paragraph("")

    # Date
    if "Date" in data and data["Date"]:
        date_text = "Nice, le " + data["Date"]
    else:
        today = datetime.today().strftime("%d %B %Y")
        date_text = "Nice, le " + today
    p_date = doc.add_paragraph(date_text)
    p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_date.paragraph_format.space_after = Pt(12)

    # Objet (centré et en gras)
    if "Objet" in data:
        p_objet = doc.add_paragraph()
        p_objet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_objet = p_objet.add_run("Objet : " + data["Objet"])
        run_objet.bold = True
        run_objet.font.size = Pt(14)
        p_objet.paragraph_format.space_after = Pt(12)

    # Salutation
    if "Salutation" in data:
        p_salutation = doc.add_paragraph(data["Salutation"])
        p_salutation.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_salutation.paragraph_format.space_after = Pt(12)

    # Corps de la lettre (en justifié)
    if "Corps" in data:
        # Remplacer les "\n" par de véritables retours à la ligne
        corps = data["Corps"].replace("\\n", "\n")
        p_corps = doc.add_paragraph(corps)
        p_corps.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_corps.paragraph_format.space_after = Pt(12)

    # Formule de politesse
    if "Formule de politesse" in data:
        p_politesse = doc.add_paragraph(data["Formule de politesse"])
        p_politesse.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_politesse.paragraph_format.space_after = Pt(12)

    # Signature
    if "Signature" in data:
        p_signature = doc.add_paragraph(data["Signature"])
        p_signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_signature.paragraph_format.space_after = Pt(12)

    # Sauvegarde du document DOCX
    doc.save(output_file)
    print(f"Lettre de motivation générée et sauvegardée dans {output_file}")


if __name__ == "__main__":
    # Fichiers CSV d'entrée
    import sys

    if len(sys.argv) == 2:
        cover_letter_csv = sys.argv[1]
    else:
        print("Usage: python cv.py <cv_csv_file>")
        sys.exit(1)

    cover_letter_output = "cover-letter-test.docx"

    # Vérification de l'existence des fichiers CSV
    if os.path.exists(cover_letter_csv):
        generate_docx_cover_letter(cover_letter_csv, cover_letter_output)
    else:
        print(f"Le fichier {cover_letter_csv} n'existe pas.")
